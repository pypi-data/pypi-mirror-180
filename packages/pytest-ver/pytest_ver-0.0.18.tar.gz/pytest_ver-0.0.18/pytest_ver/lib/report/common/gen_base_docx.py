import datetime

import docx
import docx.enum
import docx.enum.section
import docx.enum.style
import docx.enum.table
import docx.enum.text
import docx.shared
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Inches

from ... import services


# -------------------
## Base class for generating a msword docx
class GenBaseDocx:
    # -------------------
    ## constructor
    def __init__(self):
        ## holds reference to the current PDF doc
        self._doc = None

    # -------------------
    ## initialize document
    #
    # @return None
    def _doc_init(self):
        self._set_page_layout()
        self._add_styles()
        self._gen_headers_footers()

    # -------------------
    ## set page layout to landscape/portrait
    #
    # @return None
    def _set_page_layout(self):
        section = self._doc.sections[0]
        new_width, new_height = section.page_height, section.page_width
        # TODO make it landscape/portrait as requested
        section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

        # TODO compare margins to PDF margins?
        section.left_margin = docx.shared.Inches(0.5)
        section.right_margin = docx.shared.Inches(0.5)
        section.top_margin = docx.shared.Inches(1.0)
        section.bottom_margin = docx.shared.Inches(1.0)

    # -------------------
    ## add styles needed for various tables
    #
    # @return None
    def _add_styles(self):
        # Note: use keep_with_next to minimize splitting across pages

        style = self._doc.styles.add_style('ver_table1_header', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = self._doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = docx.shared.Pt(10)
        style.font.bold = True
        style.paragraph_format.space_before = docx.shared.Inches(0.08)
        style.paragraph_format.space_after = docx.shared.Inches(0.08)
        style.paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.SINGLE
        style.paragraph_format.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT
        style.paragraph_format.keep_with_next = True

        style = self._doc.styles.add_style('ver_table1_cell', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = self._doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = docx.shared.Pt(10)
        style.font.bold = False
        style.paragraph_format.space_before = docx.shared.Inches(0.0)
        style.paragraph_format.space_after = docx.shared.Inches(0.0)
        style.paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.SINGLE
        style.paragraph_format.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT
        style.paragraph_format.keep_with_next = True

        style = self._doc.styles.add_style('ver_table1_desc_cell', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = self._doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = docx.shared.Pt(10)
        style.font.bold = False
        style.paragraph_format.space_before = docx.shared.Inches(0.08)
        style.paragraph_format.space_after = docx.shared.Inches(0.0)
        style.paragraph_format.line_spacing_rule = docx.enum.text.WD_LINE_SPACING.SINGLE
        style.paragraph_format.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT
        style.paragraph_format.keep_with_next = True

    # -------------------
    ## generate test run information
    #
    # @return None
    def _gen_test_run_details(self):
        # TODO find BodyText style
        # TODO not blue title?
        self._doc.add_heading('Test Run Details', level=3)

        line = f"{'Test Run Type': <20}: {services.cfg.test_run_type}"
        self._doc.add_paragraph(line, style='List Bullet')

        line = f"{'Test Run ID': <20}: {services.cfg.test_run_id}"
        self._doc.add_paragraph(line, style='List Bullet')

        dts = datetime.datetime.now(datetime.timezone.utc).astimezone().strftime(services.cfg.dts_format)
        line = f"{'Document Generated': <20}: {dts}"
        self._doc.add_paragraph(line, style='List Bullet')

    # -------------------
    ## generate title
    #
    # @param title  the title to draw
    # @return None
    def _gen_title(self, title):
        # TODO how to not make it blue?
        self._doc.add_heading(title, level=3)

    # -------------------
    ## generate headers and footers
    #
    # @return None
    def _gen_headers_footers(self):
        self._gen_headers()
        self._gen_footers()

    # -------------------
    ## generate headers
    #
    # @return None
    def _gen_headers(self):
        header = self._doc.sections[0].header
        if services.cfg.page_info.orientation == 'portrait':
            tbl = header.add_table(rows=1, cols=3, width=Inches(7.5))
        else:
            tbl = header.add_table(rows=1, cols=3, width=Inches(10.0))
        tbl.style = 'Normal Table'
        tbl.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER

        edge_width = Inches(1.75)
        if services.cfg.page_info.orientation == 'portrait':
            # should add to 7.5"
            tbl.columns[0].width = edge_width
            tbl.columns[1].width = Inches(4.00)
            tbl.columns[2].width = edge_width
        else:
            # should add to 10"
            tbl.columns[0].width = edge_width
            tbl.columns[1].width = Inches(6.5)
            tbl.columns[2].width = edge_width

        row = tbl.rows[0].cells
        self._gen_hf_data(row, 'header')

        p = header.add_paragraph()
        self._insert_hr(p)

    # -------------------
    ## generate footers
    #
    # @return None
    def _gen_footers(self):
        footer = self._doc.sections[0].footer
        self._insert_hr(footer.paragraphs[0])
        if services.cfg.page_info.orientation == 'portrait':
            tbl = footer.add_table(rows=1, cols=3, width=Inches(7.5))
        else:
            tbl = footer.add_table(rows=1, cols=3, width=Inches(10.0))
        tbl.style = 'Normal Table'
        tbl.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER

        edge_width = Inches(1.75)
        if services.cfg.page_info.orientation == 'portrait':
            # should add to 7.5"
            tbl.columns[0].width = edge_width
            tbl.columns[1].width = Inches(4.00)
            tbl.columns[2].width = edge_width
        else:
            # should add to 10"
            tbl.columns[0].width = edge_width
            tbl.columns[1].width = Inches(6.5)
            tbl.columns[2].width = edge_width

        row = tbl.rows[0].cells
        self._gen_hf_data(row, 'footer')

    # -------------------
    ## generate headers/footer data
    #
    # @param row  the row for the header/footer data
    # @param tag  indicates header or footer
    # @return None
    def _gen_hf_data(self, row, tag):
        col = 0
        if services.cfg.page_info[tag].left == '<pageno>':
            row[col].paragraphs[0].add_run(f'Page ')  # {canvas.getPageNumber()}')
            row[col].paragraphs[0].alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            row[col].paragraphs[0].add_run(services.cfg.page_info[tag].left)
            row[col].paragraphs[0].alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT

        col = 1
        if services.cfg.page_info[tag].middle == '<pageno>':
            row[col].paragraphs[0].add_run(f'Page ')  # {canvas.getPageNumber()}')
            row[col].paragraphs[0].alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            row[col].paragraphs[0].add_run(services.cfg.page_info[tag].middle)
            row[col].paragraphs[0].alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

        col = 2
        if services.cfg.page_info[tag].right == '<pageno>':
            row[col].paragraphs[0].add_run(f'Page ')  # {canvas.getPageNumber()}')
            row[col].paragraphs[0].alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            row[col].paragraphs[0].add_run(services.cfg.page_info[tag].right)
            row[col].paragraphs[0].alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT

    # -------------------
    ## insert a header paragraph
    #
    # @param paragraph the paragraph to add to the header
    # @return None
    def _insert_hr(self, paragraph):
        p = paragraph._p  # p is the <w:p> XML element
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        pPr.insert_element_before(pBdr,
                                  'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                                  'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                                  'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                                  'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                                  'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                                  'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                                  'w:pPrChange'
                                  )
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
