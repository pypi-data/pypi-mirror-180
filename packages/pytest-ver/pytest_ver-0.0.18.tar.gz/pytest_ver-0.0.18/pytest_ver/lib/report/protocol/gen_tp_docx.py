# from docx.shared import Inches
import os

from docx import Document

from ..common.gen_base_docx import GenBaseDocx
from ... import services


# -------------------
## Generates a Test Protocol report in MSWord docx format
class GenTpDocx(GenBaseDocx):
    # -------------------
    ## constructor
    #
    # @param protocols   the data to use
    # @param do_results  generate results or not
    def __init__(self, protocols, do_results=True):
        super().__init__()
        ## holds the protocol infor
        self._protocols = protocols
        ## holds flag to generate results (for TP report), or not (for TP doc)
        self._do_results = do_results

    # -------------------
    ## generate the report
    #
    # @return None
    def gen(self):
        if self._do_results:
            services.logger.start('report: TP with results (PDF)')
        else:
            services.logger.start('report: TP without results (PDF)')

        ## see GenBaseDocx
        self._doc = Document()
        self._doc_init()

        # # TODO move into its own class
        # if services.cfg.reqmt_json_path is not None:
        #     if os.path.isfile(services.cfg.reqmt_json_path):
        #         with open(services.cfg.reqmt_json_path, 'r', encoding='utf-8') as fp:
        #             cleanj = jsmin.jsmin(fp.read())
        #             self._requirements = json.loads(cleanj)

        self._gen_test_run_details()
        if self._do_results:
            self._gen_title('Test Protocols with results')
        else:
            self._gen_title('Test Protocols')

        # for _, protocol in self._protocols.items():
        #     self._gen_protocol(protocol)

        # @@@@ junk
        # p = self._doc.add_paragraph('A plain paragraph having some ')
        # p.add_run('bold').bold = True
        # p.add_run(' and some ')
        # p.add_run('italic.').italic = True
        #
        self._doc.add_paragraph('Intense quote', style='Intense Quote')

        # TODO how to align row cells
        # TODO how to set top row a color
        records = (
            (3, '101', 'Spam'),
            (7, '422', 'Eggs'),
            (4, '631', 'Spam, spam, eggs, and spam')
        )

        table = self._doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Qty'
        hdr_cells[1].text = 'Id'
        hdr_cells[2].text = 'Desc'
        for qty, id, desc in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(qty)
            row_cells[1].text = id
            row_cells[2].text = desc
        # @@@@ junk

        self._build()

    # -------------------
    ## build the document with the current list of elements
    #
    # @return None
    def _build(self):
        if self._do_results:
            fname = services.cfg.tp_report_fname
        else:
            fname = services.cfg.tp_protocol_fname

        path = os.path.join(services.cfg.outdir, f'{fname}.docx')
        self._doc.save(path)
